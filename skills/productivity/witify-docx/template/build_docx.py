#!/usr/bin/env python3
"""Witify-branded Word document generator.

Copy this whole `template/` folder to a scratch directory, fill in the CONFIG
block, replace the CONTENT section, then run `python3 build_docx.py`.
Requires: python-docx and pymupdf. The cover is drawn straight to PDF by
pymupdf and rasterized — no browser involved. Aptos renders exactly when
Word's bundled fonts are present; otherwise the cover falls back to Helvetica.
"""
import os

import pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ============================================================= CONFIG
TITLE = "Server\nRequirements"                 # cover title ("\n" for line breaks)
SUBTITLE = "Laravel Application Hosting on Azure"
META_LINE = "Draft for discussion · August 2026"
OUT = os.path.join(os.getcwd(), "Witify Document.docx")

# ============================================================= brand
CODE_BG_HEX = "F3F4F2"
FONT = "Aptos"
MONO = "Consolas"

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")
ICON_HEADER = os.path.join(ASSETS, "witify-icon-black50.png")
COVER_PNG = os.path.join(HERE, "cover.png")
APTOS_DIR = "/Applications/Microsoft Word.app/Contents/Resources/DFonts"


# ============================================================= cover render
def render_cover():
    forest, red = (0.11, 0.169, 0.188), (1, 0.278, 0.275)          # 1C2B30, FF4746
    white, light, gray = (1, 1, 1), (0.682, 0.741, 0.761), (0.561, 0.627, 0.651)

    cover = pymupdf.open()
    page = cover.new_page(width=850, height=1100)
    page.draw_rect(page.rect, fill=forest, color=None)

    def svg(name, rect):
        src = pymupdf.open("pdf", pymupdf.open(os.path.join(ASSETS, name)).convert_to_pdf())
        page.show_pdf_page(rect, src, 0)

    svg("witify-icon-tonal.svg", pymupdf.Rect(300, 722.5, 1060, 1220))  # watermark, bleeds off-page
    svg("logo-light.svg", pymupdf.Rect(72, 64, 240, 106.3))
    page.draw_rect(pymupdf.Rect(72, 440, 136, 444), fill=red, color=None)

    def text(pos, s, size, color, bold=False):
        ttf = os.path.join(APTOS_DIR, "Aptos-Bold.ttf" if bold else "Aptos.ttf")
        if os.path.isfile(ttf):
            page.insert_text(pos, s, fontsize=size, color=color, fontfile=ttf,
                             fontname="AptosBold" if bold else "Aptos")
        else:
            page.insert_text(pos, s, fontsize=size, color=color,
                             fontname="hebo" if bold else "helv")

    y = 480
    for line in TITLE.splitlines():                                # 58px title, 1.06 line-height
        text((72, y + 45.5), line, 58, white, bold=True)
        y += 61.5
    text((72, y + 43), SUBTITLE, 20, light)
    text((72, 996), META_LINE, 13, white, bold=True)
    text((72, 1019.5), "Witify Technologies · witify.io", 13, gray)

    page.get_pixmap(dpi=144).save(COVER_PNG)                       # 2x the 850x1100 design


render_cover()

# ============================================================= document setup
doc = Document()

MARGIN = 0.75
section = doc.sections[0]
section.top_margin = Inches(MARGIN)
section.bottom_margin = Inches(MARGIN)
section.left_margin = Inches(MARGIN)
section.right_margin = Inches(MARGIN)


def style_font(style, name=FONT, size=None, bold=None):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    # Theme font attributes beat explicit names in Word; strip them.
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    # Automatic (black) color, overriding built-in heading blues.
    existing_color = rpr.find(qn("w:color"))
    if existing_color is not None:
        rpr.remove(existing_color)
    auto = OxmlElement("w:color")
    auto.set(qn("w:val"), "auto")
    rpr.append(auto)


styles = doc.styles
style_font(styles["Normal"], size=11)
style_font(styles["Title"], size=24, bold=True)
style_font(styles["Heading 1"], size=16, bold=True)
style_font(styles["Heading 2"], size=13, bold=True)
style_font(styles["Heading 3"], size=11.5, bold=True)
style_font(styles["List Bullet"], size=11)
style_font(styles["List Number"], size=11)

styles["Heading 1"].paragraph_format.space_after = Pt(10)
styles["Heading 2"].paragraph_format.space_after = Pt(6)

# ============================================================= helpers
def set_run_font(run, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def add_run(p, text, size=None, bold=False, italic=False, color=None, font=FONT):
    run = p.add_run(text)
    set_run_font(run, font)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def anchor_picture_to_page(inline_shape, width_emu, height_emu):
    """Convert an inline picture to a floating one anchored at the page's top-left."""
    inline = inline_shape._inline
    anchor = OxmlElement("wp:anchor")
    for key, value in (("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
                       ("simplePos", "0"), ("relativeHeight", "0"), ("behindDoc", "1"),
                       ("locked", "0"), ("layoutInCell", "1"), ("allowOverlap", "1")):
        anchor.set(key, value)
    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)
    for tag in ("wp:positionH", "wp:positionV"):
        pos = OxmlElement(tag)
        pos.set("relativeFrom", "page")
        offset = OxmlElement("wp:posOffset")
        offset.text = "0"
        pos.append(offset)
        anchor.append(pos)
    extent = inline.find(qn("wp:extent"))
    extent.set("cx", str(width_emu))
    extent.set("cy", str(height_emu))
    anchor.append(extent)
    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(inline.find(qn("wp:docPr")))
    frame_pr = inline.find(qn("wp:cNvGraphicFramePr"))
    if frame_pr is not None:
        anchor.append(frame_pr)
    anchor.append(inline.find(qn("a:graphic")))
    drawing = inline.getparent()
    drawing.replace(inline, anchor)


def shade_cell(cell, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def no_row_split(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))


def cell_para(cell, first=True):
    p = cell.paragraphs[0] if first and not cell.paragraphs[0].runs else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def bullet(parts, style="List Bullet"):
    """parts: list of (text, opts) tuples, e.g. [("Bold lead", B), (" rest.", N)]."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    for text, opts in parts:
        add_run(p, text, **opts)
    return p


def styled_table(headers, rows, widths):
    """rows: list of row lists; each cell is a string or (main, sub_gray_line)."""
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    table.autofit = False
    for cell, label in zip(table.rows[0].cells, headers):
        add_run(cell_para(cell), label, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            main, sub = value if isinstance(value, tuple) else (value, None)
            add_run(cell_para(cells[i]), main, bold=(i == 0), size=10)
            if sub:
                sp = cells[i].add_paragraph()
                sp.paragraph_format.space_before = Pt(0)
                sp.paragraph_format.space_after = Pt(2)
                add_run(sp, sub, size=9.5, color=RGBColor(0x5E, 0x6A, 0x6E))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
    no_row_split(table)
    return table


def code_block(lines):
    table = doc.add_table(rows=1, cols=1, style="Table Grid")
    no_row_split(table)
    cell = table.rows[0].cells[0]
    shade_cell(cell, CODE_BG_HEX)
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_run(p, line if line else " ", size=9, font=MONO)
    return table


B = {"bold": True}
N = {}
CODE_INLINE = {"font": MONO, "size": 10}

# ============================================================= cover page
EMU_PER_INCH = 914400

cover_p = doc.add_paragraph()
cover_p.paragraph_format.space_after = Pt(0)
cover_shape = cover_p.add_run().add_picture(COVER_PNG, width=Inches(8.5))
anchor_picture_to_page(cover_shape,
                       int(8.5 * EMU_PER_INCH),
                       int(11 * EMU_PER_INCH))
doc.add_page_break()

# ============================================================= CONTENT
# Replace everything in this section with the document's actual content.
# Available: doc.add_heading("1. Title", level=1), doc.add_paragraph("..."),
# bullet([...]), styled_table(...), code_block([...]), add_run for rich runs.

doc.add_heading("1. Example section", level=1)
doc.add_paragraph("An intro paragraph under a Heading 1 section.")

doc.add_heading("Example subsection", level=2)
bullet([("Bold lead", B), (" followed by normal text and ", N), ("inline code", CODE_INLINE), (".", N)])
bullet([("A second bullet.", N)])

styled_table(
    headers=("Component", "Requirement"),
    rows=[
        ("OS", ("Ubuntu LTS", "A gray explanatory sub-line.")),
        ("Web server", "Nginx"),
    ],
    widths=(Inches(1.5), Inches(5.5)),
)

doc.add_heading("Example script", level=2)
code_block([
    "echo 'hello'",
    "# a comment",
])

# ============================================================= header, footer
section.different_first_page_header_footer = True
section.header_distance = Inches(0.4)

# Content pages number from 1 (the cover is page 0 and shows no number).
pg_num_type = OxmlElement("w:pgNumType")
pg_num_type.set(qn("w:start"), "0")
section._sectPr.append(pg_num_type)

# Header: Witify icon on the right.
icon_p = section.header.paragraphs[0]
icon_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
icon_p.paragraph_format.space_before = Pt(0)
icon_p.paragraph_format.space_after = Pt(0)
icon_p.add_run().add_picture(ICON_HEADER, height=Inches(0.2))

# Footer: brand on the left, page number on the right.
footer_p = section.footer.paragraphs[0]
footer_tabs = OxmlElement("w:tabs")
for pos in ("4680", "9360"):
    clear_tab = OxmlElement("w:tab")
    clear_tab.set(qn("w:val"), "clear")
    clear_tab.set(qn("w:pos"), pos)
    footer_tabs.append(clear_tab)
right_tab = OxmlElement("w:tab")
right_tab.set(qn("w:val"), "right")
right_tab.set(qn("w:pos"), "10080")
footer_tabs.append(right_tab)
footer_p.paragraph_format.element.get_or_add_pPr().append(footer_tabs)
add_run(footer_p, "Witify Technologies · witify.io", size=8, color=RGBColor(0x5E, 0x6A, 0x6E))
add_run(footer_p, "\t", size=8)

page_run = add_run(footer_p, "", size=8, color=RGBColor(0x5E, 0x6A, 0x6E))
for el_name, attrs, text in (
    ("w:fldChar", {"w:fldCharType": "begin"}, None),
    ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
    ("w:fldChar", {"w:fldCharType": "separate"}, None),
    ("w:t", {}, "1"),
    ("w:fldChar", {"w:fldCharType": "end"}, None),
):
    el = OxmlElement(el_name)
    for attr, value in attrs.items():
        el.set(qn(attr), value)
    if text is not None:
        el.text = text
    page_run._element.append(el)

doc.save(OUT)
print(f"Saved: {OUT}")
